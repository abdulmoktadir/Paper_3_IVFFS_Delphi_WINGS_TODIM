import io
import base64
import math
import numpy as np
import pandas as pd
import streamlit as st

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# graphviz is optional on some deployments (binary may be missing)
try:
    import graphviz
    GRAPHVIZ_OK = True
except Exception:
    GRAPHVIZ_OK = False


# =========================================================
# Helpers: Safe display formatting (FIXES YOUR ERROR)
# =========================================================
def safe_numeric_style(df: pd.DataFrame, precision: int = 6):
    """
    Streamlit sometimes crashes when Styler.format is applied to mixed types.
    This formats ONLY numeric columns; leaves others unchanged.
    """
    df2 = df.copy()
    num_cols = df2.select_dtypes(include=[np.number]).columns
    # avoid Styler on huge tables if you want; but we keep it safe
    sty = df2.style
    if len(num_cols) > 0:
        fmt = "{:." + str(precision) + "f}"
        sty = sty.format({c: fmt for c in num_cols})
    return sty


# =========================================================
# IVFFS representation
# We treat an IVFF number as (mu_L, mu_U, nu_L, nu_U)
# =========================================================
IVFF = tuple[float, float, float, float]


def clamp01(x: float) -> float:
    return max(0.0, min(1.0, float(x)))


def make_ivff(mu_l, mu_u, nu_l, nu_u) -> IVFF:
    mu_l, mu_u, nu_l, nu_u = map(clamp01, [mu_l, mu_u, nu_l, nu_u])
    if mu_l > mu_u:
        mu_l, mu_u = mu_u, mu_l
    if nu_l > nu_u:
        nu_l, nu_u = nu_u, nu_l
    return (mu_l, mu_u, nu_l, nu_u)


def ivff_to_str(x: IVFF, nd: int = 4) -> str:
    mu_l, mu_u, nu_l, nu_u = x
    return f"([{{:.{nd}f}}, {{:.{nd}f}}], [{{:.{nd}f}}, {{:.{nd}f}}])".format(mu_l, mu_u, nu_l, nu_u)


def parse_ivff_str(s: str) -> IVFF | None:
    """
    Parses strings like: ([0.05,0.15],[0.85,0.95])
    """
    try:
        t = (
            s.strip()
            .replace("(", "")
            .replace(")", "")
            .replace("[", "")
            .replace("]", "")
            .replace(" ", "")
        )
        parts = t.split(",")
        if len(parts) != 4:
            return None
        a, b, c, d = map(float, parts)
        return make_ivff(a, b, c, d)
    except Exception:
        return None


# =========================================================
# IVFFS crisp functions
# (Fermatean-style score; works for ranking/defuzz)
# =========================================================
def score_ivff(x: IVFF) -> float:
    """
    A practical IVFFS score (Fermatean-style):
      S = avg(mu^3) - avg(nu^3)
    Range approx [-1, 1]
    """
    mu_l, mu_u, nu_l, nu_u = x
    mu = (mu_l**3 + mu_u**3) / 2.0
    nu = (nu_l**3 + nu_u**3) / 2.0
    return mu - nu


def exp_value_ivff(x: IVFF) -> float:
    """
    Expected value proxy used for DEMATEL/WINGS crisp matrix:
      EV = avg(mu)
    Range [0,1]
    """
    mu_l, mu_u, _, _ = x
    return (mu_l + mu_u) / 2.0


# =========================================================
# IVFFS aggregation (expert aggregation)
# =========================================================
def ivff_weighted_aggregate(ivffs: list[IVFF], weights: list[float]) -> IVFF:
    """
    Robust weighted aggregation for IVFF:
      mu = (sum w * mu^3)^(1/3)
      nu = (sum w * nu^3)^(1/3)
    applied to lower/upper separately.
    """
    if len(ivffs) == 0:
        return make_ivff(0, 0, 0, 0)

    w = np.array(weights, dtype=float)
    if w.sum() == 0:
        w = np.ones_like(w) / len(w)
    else:
        w = w / w.sum()

    mu_l = (np.sum(w * np.array([x[0] ** 3 for x in ivffs])) ) ** (1/3)
    mu_u = (np.sum(w * np.array([x[1] ** 3 for x in ivffs])) ) ** (1/3)
    nu_l = (np.sum(w * np.array([x[2] ** 3 for x in ivffs])) ) ** (1/3)
    nu_u = (np.sum(w * np.array([x[3] ** 3 for x in ivffs])) ) ** (1/3)

    return make_ivff(mu_l, mu_u, nu_l, nu_u)


# =========================================================
# Load linguistic scales from your Excel template
# (sheet: "Experts Data" in your uploaded workbook)
# =========================================================
def load_scales_from_excel(xlsx_file) -> dict:
    """
    Expected layout (matches your template):
      - Influence assessment table around T3 / U4
      - Strength (relevance) table around U18
      - TODIM linguistic table around U4 (same area in many templates)

    We detect by scanning for abbreviations + numeric columns a,b,c,d.
    """
    import openpyxl

    wb = openpyxl.load_workbook(xlsx_file, data_only=False)
    if "Experts Data" not in wb.sheetnames:
        raise ValueError("Excel must contain a sheet named 'Experts Data'.")

    ws = wb["Experts Data"]

    # Read a large rectangular area and locate tables by headers
    # We'll just scan rows to find any row that contains: ["a","b","c","d"]
    max_row = min(ws.max_row, 400)
    max_col = min(ws.max_column, 80)

    rows = []
    for r in range(1, max_row + 1):
        row = []
        for c in range(1, max_col + 1):
            row.append(ws.cell(r, c).value)
        rows.append(row)

    df = pd.DataFrame(rows)

    def find_header_row(header_tokens=("a", "b", "c", "d")):
        for i in range(len(df)):
            row = df.iloc[i].astype(str).str.lower().tolist()
            if all(tok in row for tok in header_tokens):
                return i
        return None

    # Influence table: usually has "Influence assessement" and "ELI, VLI..."
    # Strength table: usually has "Strength/relevance assessement" and "VLR..."
    # TODIM table: usually has "Linguistic assessment" and letters (a,b,c,d)

    influence_map = {}
    strength_map = {}
    todim_map = {}

    # Scan for known abbreviations and pick their numeric columns
    def grab_table_by_abbr(abbr_list):
        mapping = {}
        for i in range(len(df)):
            row = df.iloc[i].tolist()
            for j, val in enumerate(row):
                if val in abbr_list:
                    # try read a,b,c,d in next columns (common)
                    # find numeric candidates in row
                    nums = []
                    for k in range(j + 1, min(j + 10, df.shape[1])):
                        v = row[k]
                        if isinstance(v, (int, float)) and not pd.isna(v):
                            nums.append(float(v))
                    if len(nums) >= 4:
                        mapping[str(val)] = make_ivff(nums[0], nums[1], nums[2], nums[3])
        return mapping

    influence_map = grab_table_by_abbr(["ELI", "VLI", "LI", "MI", "HI", "VHI", "EHI"])
    strength_map = grab_table_by_abbr(["VLR", "LR", "MR", "HR", "VHR"])

    # TODIM linguistic terms in your earlier snippets (example: "MG", "G"...)
    # Your template may include different ones; we also capture generic 2-letter codes.
    todim_map = grab_table_by_abbr(["VP", "P", "MP", "F", "MG", "G", "VG", "VB", "MB", "M", "EB", "EG"])

    return {
        "influence": influence_map,
        "strength": strength_map,
        "todim": todim_map,
    }


# =========================================================
# WINGS (crisp DEMATEL-like core, with IVFFS inputs)
# =========================================================
def wings_compute(
    component_names: list[str],
    strengths_by_expert: list[list[IVFF]],
    influence_by_expert: list[list[list[IVFF]]],
    expert_weights: list[float],
):
    n = len(component_names)
    m = len(strengths_by_expert)

    # Aggregate expert IVFFS matrices
    agg_strength = []
    for i in range(n):
        ivffs = [strengths_by_expert[e][i] for e in range(m)]
        agg_strength.append(ivff_weighted_aggregate(ivffs, expert_weights))

    agg_influence = [[make_ivff(0, 0, 0, 0) for _ in range(n)] for _ in range(n)]
    for i in range(n):
        for j in range(n):
            if i == j:
                agg_influence[i][j] = agg_strength[i]
            else:
                ivffs = [influence_by_expert[e][i][j] for e in range(m)]
                agg_influence[i][j] = ivff_weighted_aggregate(ivffs, expert_weights)

    # Build crisp direct matrix D using expected value of IVFF
    D = np.zeros((n, n), dtype=float)
    for i in range(n):
        for j in range(n):
            D[i, j] = exp_value_ivff(agg_influence[i][j])

    # Normalize by max row sum (DEMATEL normalization)
    row_sums = D.sum(axis=1)
    s = row_sums.max() if row_sums.max() != 0 else 1.0
    X = D / s

    # Total relation matrix: T = X (I - X)^-1
    I = np.eye(n)
    try:
        T = X @ np.linalg.inv(I - X)
    except np.linalg.LinAlgError:
        # fallback pseudo-inverse
        T = X @ np.linalg.pinv(I - X)

    TI = T.sum(axis=1)
    TR = T.sum(axis=0)
    engagement = TI + TR
    role = TI - TR

    # Weights from engagement (safe, stable)
    w = engagement.copy()
    w = np.maximum(w, 0.0)
    w_sum = w.sum() if w.sum() != 0 else 1.0
    weights = w / w_sum

    out = pd.DataFrame({
        "Component": component_names,
        "TI": TI,
        "TR": TR,
        "Engagement": engagement,
        "Role": role,
        "Expected value": engagement,   # shown explicitly for your table naming
        "Weight": weights
    })

    return {
        "agg_strength": agg_strength,
        "agg_influence": agg_influence,
        "D": D,
        "X": X,
        "T": T,
        "out": out
    }


# =========================================================
# TODIM (crisp TODIM using IVFFS score)
# =========================================================
def todim_rank(
    decision_scores: np.ndarray,   # m x n (normalized to [0,1])
    weights: np.ndarray,           # n
    theta: float = 1.0,
    alpha: float = 1.0
):
    """
    Standard TODIM-like dominance.
    Uses reference criterion r = argmax(weights).
    """
    m, n = decision_scores.shape
    w = weights.astype(float)
    if w.sum() == 0:
        w = np.ones(n) / n
    else:
        w = w / w.sum()

    r = int(np.argmax(w))
    wr = w[r] if w[r] != 0 else 1e-9

    delta = np.zeros((m, m), dtype=float)

    for i in range(m):
        for k in range(m):
            if i == k:
                continue
            s = 0.0
            for j in range(n):
                diff = decision_scores[i, j] - decision_scores[k, j]
                if diff >= 0:
                    s += (w[j] / wr) * (diff ** alpha)
                else:
                    # loss part
                    wj = w[j] if w[j] != 0 else 1e-9
                    s -= (1.0 / theta) * (wr / wj) * ((-diff) ** alpha)
            delta[i, k] = s

    phi = delta.sum(axis=1)
    # normalize to [0,1]
    mn, mx = phi.min(), phi.max()
    if mx - mn == 0:
        xi = np.zeros_like(phi)
    else:
        xi = (phi - mn) / (mx - mn)

    return phi, xi


# =========================================================
# Word report export (WINGS)
# =========================================================
def make_word_report_wings(out_df: pd.DataFrame) -> Document:
    doc = Document()
    title = doc.add_heading("IVFFS–WINGS Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Results summary (TI / TR / Engagement / Role / Weight).")

    table = doc.add_table(rows=1, cols=len(out_df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(out_df.columns):
        hdr[i].text = str(c)

    for _, row in out_df.iterrows():
        r = table.add_row().cells
        for i, c in enumerate(out_df.columns):
            v = row[c]
            if isinstance(v, (float, int, np.floating, np.integer)):
                r[i].text = f"{float(v):.6f}"
            else:
                r[i].text = str(v)

    return doc


def word_download_link(doc: Document, filename="IVFFS_WINGS_Report.docx"):
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    b64 = base64.b64encode(bio.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Word Report</a>'
    return href


# =========================================================
# UI: IVFFS–WINGS module
# =========================================================
def ivffs_wings_module():
    st.header("🧩 IVFFS–WINGS")

    st.info("Upload your Excel template to auto-load linguistic scales (recommended).")

    up = st.file_uploader("Upload Excel (your template)", type=["xlsx"], key="ivffs_excel")
    if up is None:
        st.warning("Please upload your Excel template to proceed.")
        return

    try:
        scales = load_scales_from_excel(up)
    except Exception as e:
        st.error(f"Could not read scales from Excel: {e}")
        return

    strength_scale = scales["strength"]
    influence_scale = scales["influence"]

    if len(strength_scale) == 0 or len(influence_scale) == 0:
        st.error("Scale tables not detected in Excel (Strength/Influence). Please ensure sheet 'Experts Data' contains them.")
        return

    with st.expander("Loaded Linguistic Scales", expanded=False):
        st.write("**Strength terms**")
        st.dataframe(pd.DataFrame(
            [{"abbr": k, "ivff": ivff_to_str(v, 4)} for k, v in strength_scale.items()]
        ), hide_index=True, use_container_width=True)
        st.write("**Influence terms**")
        st.dataframe(pd.DataFrame(
            [{"abbr": k, "ivff": ivff_to_str(v, 4)} for k, v in influence_scale.items()]
        ), hide_index=True, use_container_width=True)

    # config
    st.subheader("Configuration")
    n_components = st.number_input("Number of components", min_value=2, max_value=40, value=5, step=1, key="w_ncomp_ivffs")
    n_experts = st.number_input("Number of experts", min_value=1, max_value=30, value=2, step=1, key="w_nexp_ivffs")

    component_names = []
    cols = st.columns(min(5, n_components))
    for i in range(n_components):
        with cols[i % len(cols)]:
            component_names.append(st.text_input(f"Component {i+1}", value=f"C{i+1}", key=f"w_comp_ivffs_{i}"))

    st.markdown("### Expert weights")
    if n_experts == 1:
        expert_weights = [1.0]
        st.success("Single expert → weight = 1.0")
    else:
        expert_weights = []
        ecols = st.columns(min(6, n_experts))
        for e in range(n_experts):
            with ecols[e % len(ecols)]:
                expert_weights.append(
                    st.number_input(
                        f"E{e+1}",
                        min_value=0.0, max_value=1.0,
                        value=round(1 / n_experts, 5),
                        step=0.00001,
                        format="%.5f",
                        key=f"w_expw_ivffs_{e}"
                    )
                )
        if not np.isclose(sum(expert_weights), 1.0):
            st.error(f"Expert weights must sum to 1.0 (now: {sum(expert_weights):.6f}).")
            return

    # editors per expert
    st.subheader("Expert inputs")

    tabs = st.tabs([f"Expert {e+1}" for e in range(n_experts)])
    strengths_by_expert = []
    influence_by_expert = []

    strength_terms = list(strength_scale.keys())
    influence_terms = list(influence_scale.keys())

    for e in range(n_experts):
        with tabs[e]:
            st.markdown("#### Strength (diagonal)")
            srow = []
            scol = st.columns(min(6, n_components))
            for i in range(n_components):
                with scol[i % len(scol)]:
                    term = st.selectbox(
                        f"{component_names[i]} strength",
                        options=strength_terms,
                        index=strength_terms.index("HR") if "HR" in strength_terms else 0,
                        key=f"w_strength_{e}_{i}"
                    )
                    srow.append(strength_scale[term])
            strengths_by_expert.append(srow)

            st.markdown("#### Influence matrix (row influences column)")
            mat = [[make_ivff(0, 0, 0, 0) for _ in range(n_components)] for _ in range(n_components)]
            for i in range(n_components):
                row_cols = st.columns(min(6, n_components))
                for j in range(n_components):
                    with row_cols[j % len(row_cols)]:
                        if i == j:
                            st.markdown("—")
                            mat[i][j] = strength_scale.get("HR", srow[i])
                        else:
                            term = st.selectbox(
                                f"{component_names[i]}→{component_names[j]}",
                                options=influence_terms,
                                index=influence_terms.index("ELI") if "ELI" in influence_terms else 0,
                                key=f"w_infl_{e}_{i}_{j}"
                            )
                            mat[i][j] = influence_scale[term]
            influence_by_expert.append(mat)

    if st.button("🚀 Run IVFFS–WINGS", type="primary", use_container_width=True, key="run_ivffs_wings"):
        with st.spinner("Computing WINGS..."):
            res = wings_compute(component_names, strengths_by_expert, influence_by_expert, expert_weights)

        st.success("Done.")

        out = res["out"]

        st.subheader("Results table (TI / TR / Engagement / Role / Expected value / Weight)")
        # FIX: Use safe_numeric_style (prevents your ValueError crash)
        st.dataframe(safe_numeric_style(out, precision=6), use_container_width=True, hide_index=True)

        with st.expander("Matrices (crisp)"):
            st.write("Direct matrix D (expected values)")
            st.dataframe(pd.DataFrame(res["D"], index=component_names, columns=component_names).round(6), use_container_width=True)
            st.write("Normalized matrix X")
            st.dataframe(pd.DataFrame(res["X"], index=component_names, columns=component_names).round(6), use_container_width=True)
            st.write("Total relation matrix T")
            st.dataframe(pd.DataFrame(res["T"], index=component_names, columns=component_names).round(6), use_container_width=True)

        st.subheader("Export")
        doc = make_word_report_wings(out)
        st.markdown(word_download_link(doc), unsafe_allow_html=True)

        # store weights for TODIM module
        st.session_state["wings_weights"] = out[["Component", "Weight"]].copy()


# =========================================================
# UI: IVFFS–TODIM module
# =========================================================
def ivffs_todim_module():
    st.header("📌 IVFFS–TODIM")

    st.info("If you ran IVFFS–WINGS already, TODIM can auto-load criterion weights from WINGS output.")

    up = st.file_uploader("Upload Excel (same template, to load TODIM linguistic scale)", type=["xlsx"], key="todim_excel")
    if up is None:
        st.warning("Upload the Excel template to proceed.")
        return

    try:
        scales = load_scales_from_excel(up)
    except Exception as e:
        st.error(f"Could not read scales from Excel: {e}")
        return

    todim_scale = scales["todim"]
    if len(todim_scale) == 0:
        # fallback: reuse influence if needed
        todim_scale = scales["influence"]

    st.subheader("Define Alternatives and Criteria")
    c1, c2 = st.columns(2)
    alts_in = c1.text_input("Alternatives (comma-separated)", "A1, A2, A3", key="todim_alts")
    crits_in = c2.text_input("Criteria (comma-separated)", "C1, C2, C3", key="todim_crits")

    alternatives = [a.strip() for a in alts_in.split(",") if a.strip()]
    criteria = [c.strip() for c in crits_in.split(",") if c.strip()]
    if len(alternatives) < 2 or len(criteria) < 1:
        st.warning("Provide at least 2 alternatives and 1 criterion.")
        return

    # Weights table (allow 5 digits + smaller step)
    st.markdown("### Criteria types and weights")

    # if WINGS weights exist, auto-fill
    wings_w = st.session_state.get("wings_weights", None)
    default_w = [round(1 / len(criteria), 5)] * len(criteria)
    if len(criteria) > 0:
        default_w[-1] = round(1.0 - sum(default_w[:-1]), 5)

    if wings_w is not None:
        # map by name if possible
        w_map = dict(zip(wings_w["Component"], wings_w["Weight"]))
        default_w = [float(w_map.get(c, default_w[i])) for i, c in enumerate(criteria)]
        # re-normalize
        s = sum(default_w)
        default_w = [x / s for x in default_w] if s != 0 else default_w

    if "todim_crit_df" not in st.session_state or set(st.session_state["todim_crit_df"]["Criterion"]) != set(criteria):
        st.session_state["todim_crit_df"] = pd.DataFrame({
            "Criterion": criteria,
            "Type": ["Benefit"] * len(criteria),
            "Weight": default_w
        })

    edited = st.data_editor(
        st.session_state["todim_crit_df"],
        hide_index=True,
        use_container_width=True,
        column_config={
            "Type": st.column_config.SelectboxColumn("Type", options=["Benefit", "Cost"]),
            # FIX: allow 5 digits + tiny step
            "Weight": st.column_config.NumberColumn("Weight", format="%.5f", min_value=0.0, max_value=1.0, step=0.00001),
        },
        key="todim_crit_editor"
    )

    w = edited["Weight"].astype(float).to_numpy()
    if not np.isclose(w.sum(), 1.0):
        st.error(f"Criteria weights must sum to 1.0 (now: {w.sum():.6f}).")
        return

    types = edited["Type"].tolist()

    st.subheader("Experts and evaluations")
    n_experts = st.number_input("Number of experts", min_value=1, max_value=30, value=2, step=1, key="todim_nexp")

    st.markdown("**Expert weights** (sum to 1)")
    if n_experts == 1:
        ew = [1.0]
        st.success("Single expert → weight = 1.0")
    else:
        ew = []
        cols = st.columns(min(6, n_experts))
        for i in range(n_experts):
            with cols[i % len(cols)]:
                ew.append(
                    st.number_input(
                        f"E{i+1}",
                        min_value=0.0, max_value=1.0,
                        value=round(1 / n_experts, 5),
                        step=0.00001,
                        format="%.5f",
                        key=f"todim_ew_{i}"
                    )
                )
        if not np.isclose(sum(ew), 1.0):
            st.error(f"Expert weights must sum to 1.0 (now: {sum(ew):.6f}).")
            return

    term_list = list(todim_scale.keys())
    with st.expander("TODIM linguistic scale loaded", expanded=False):
        st.dataframe(
            pd.DataFrame([{"abbr": k, "ivff": ivff_to_str(v, 4)} for k, v in todim_scale.items()]),
            hide_index=True,
            use_container_width=True
        )

    # editors per expert: alternatives x criteria linguistic
    if "todim_expert_tables" not in st.session_state:
        st.session_state["todim_expert_tables"] = {}

    need_reset = (
        len(st.session_state["todim_expert_tables"]) != n_experts
        or (n_experts > 0 and (
            set(st.session_state["todim_expert_tables"].get(0, pd.DataFrame()).index) != set(alternatives)
            or set(st.session_state["todim_expert_tables"].get(0, pd.DataFrame()).columns) != set(criteria)
        ))
    )
    if need_reset:
        st.session_state["todim_expert_tables"] = {
            e: pd.DataFrame(term_list[0], index=alternatives, columns=criteria) for e in range(n_experts)
        }

    tabs = st.tabs([f"Expert {e+1}" for e in range(n_experts)])
    for e in range(n_experts):
        with tabs[e]:
            st.session_state["todim_expert_tables"][e] = st.data_editor(
                st.session_state["todim_expert_tables"][e],
                use_container_width=True,
                column_config={
                    c: st.column_config.SelectboxColumn(c, options=term_list) for c in criteria
                },
                key=f"todim_editor_{e}"
            )

    st.subheader("Run TODIM")
    theta = st.number_input("θ (theta) loss attenuation", min_value=0.01, max_value=20.0, value=1.0, step=0.1, key="todim_theta")
    alpha = st.number_input("α (alpha) gain/loss exponent", min_value=0.1, max_value=5.0, value=1.0, step=0.1, key="todim_alpha")

    if st.button("✅ Run IVFFS–TODIM", type="primary", use_container_width=True, key="run_todim"):
        with st.spinner("Computing TODIM..."):
            # Aggregate decision matrix in IVFF domain
            agg = [[None for _ in criteria] for _ in alternatives]
            for i, alt in enumerate(alternatives):
                for j, crit in enumerate(criteria):
                    ivffs = []
                    for e in range(n_experts):
                        term = st.session_state["todim_expert_tables"][e].loc[alt, crit]
                        ivffs.append(todim_scale[term])
                    agg[i][j] = ivff_weighted_aggregate(ivffs, ew)

            agg_df = pd.DataFrame(
                [[ivff_to_str(agg[i][j], 4) for j in range(len(criteria))] for i in range(len(alternatives))],
                index=alternatives, columns=criteria
            )
            st.markdown("#### Aggregated IVFF decision matrix")
            st.dataframe(agg_df, use_container_width=True)

            # Defuzzify to crisp scores for TODIM
            # Use score shifted to [0,1]
            crisp = np.zeros((len(alternatives), len(criteria)), dtype=float)
            for i in range(len(alternatives)):
                for j in range(len(criteria)):
                    s = score_ivff(agg[i][j])      # [-1,1] approx
                    crisp[i, j] = (s + 1.0) / 2.0  # -> [0,1]

            crisp_df = pd.DataFrame(crisp, index=alternatives, columns=criteria)
            st.markdown("#### Crisp matrix (score shifted to [0,1])")
            st.dataframe(crisp_df.round(6), use_container_width=True)

            # Normalize by type (Benefit/Cost) to [0,1]
            norm = crisp.copy()
            for j in range(len(criteria)):
                col = norm[:, j]
                mn, mx = float(np.min(col)), float(np.max(col))
                if mx - mn == 0:
                    norm[:, j] = 0.0
                else:
                    if str(types[j]).lower().startswith("b"):
                        norm[:, j] = (col - mn) / (mx - mn)
                    else:
                        norm[:, j] = (mx - col) / (mx - mn)

            norm_df = pd.DataFrame(norm, index=alternatives, columns=criteria)
            st.markdown("#### Normalized matrix for TODIM")
            st.dataframe(norm_df.round(6), use_container_width=True)

            phi, xi = todim_rank(norm, w, theta=theta, alpha=alpha)

            res_df = pd.DataFrame({
                "Alternative": alternatives,
                "Phi (dominance sum)": phi,
                "Xi (normalized value)": xi
            }).sort_values("Xi (normalized value)", ascending=False).reset_index(drop=True)

            res_df["Rank"] = np.arange(1, len(res_df) + 1)

            st.markdown("#### TODIM final ranking")
            st.dataframe(safe_numeric_style(res_df, precision=6), use_container_width=True, hide_index=True)


# =========================================================
# Main app navigation
# =========================================================
def main():
    st.set_page_config(page_title="IVFFS Toolkit (WINGS + TODIM)", layout="wide")
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Choose module", ["IVFFS–WINGS", "IVFFS–TODIM"])

    if page == "IVFFS–WINGS":
        ivffs_wings_module()
    else:
        ivffs_todim_module()


if __name__ == "__main__":
    main()
